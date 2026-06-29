"""
Utility functions for processing Word documents containing questions
Supports both English and Vietnamese formats
With validation, error reporting, and image extraction
"""
import docx
import re
import io
from concurrent.futures import ThreadPoolExecutor, as_completed
from django.db import transaction
from question_bank.models import Question, AnswerOption
from question_bank.duplicate_utils import build_existing_fingerprint_set, compute_fingerprint

try:
    from docx.oxml.ns import qn
except ImportError:
    qn = None

try:
    import cloudinary.uploader
    HAS_CLOUDINARY = True
except ImportError:
    HAS_CLOUDINARY = False


MAX_FILE_SIZE_MB = 10
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024


def upload_image_to_cloudinary(image_blob, content_type='image/png'):
    """Upload image bytes to Cloudinary and return URL."""
    if not HAS_CLOUDINARY:
        print("[Cloudinary] cloudinary package not available, skipping image upload")
        return None
    try:
        result = cloudinary.uploader.upload(
            io.BytesIO(image_blob),
            folder="question_bank",
            resource_type="image",
        )
        return result.get('secure_url')
    except Exception as e:
        print(f"[Cloudinary] Upload failed: {e}")
        return None


class WordQuestionParser:
    """
    Parser for Word documents containing questions.
    Supports English and Vietnamese formats with image extraction.
    """
    
    # English patterns
    QUESTION_PATTERN_EN = re.compile(r'^Question\s+(\d+):\s*(.+)$', re.IGNORECASE)
    ANSWER_PATTERN_EN = re.compile(r'^ANSWER:\s*(.+)$', re.IGNORECASE)
    ORDER_PATTERN_EN = re.compile(r'^CORRECT\s*ORDER:\s*(.+)$', re.IGNORECASE)
    
    # Vietnamese patterns
    QUESTION_PATTERN_VI = re.compile(r'^Câu\s+(\d+):\s*(.+)$', re.IGNORECASE)
    ANSWER_PATTERN_VI = re.compile(r'^ĐÁP\s*ÁN:\s*(.+)$', re.IGNORECASE)
    ORDER_PATTERN_VI = re.compile(r'^THỨ\s*TỰ\s*ĐÚNG:\s*(.+)$', re.IGNORECASE)
    
    # Common patterns
    OPTION_PATTERN = re.compile(r'^([A-Z])\.\s*(.+)$')
    NUMBERED_OPTION_PATTERN = re.compile(r'^(\d+)\.\s*(.+)$')
    
    # Keywords for question type detection
    ORDERING_KEYWORDS_EN = ['arrange', 'order', 'sequence', 'sort']
    ORDERING_KEYWORDS_VI = ['sắp xếp', 'thứ tự']
    
    TRUE_FALSE_KEYWORDS_EN = ['true/false', 'true or false', 'evaluate', 't/f']
    TRUE_FALSE_KEYWORDS_VI = ['đúng/sai', 'đánh giá']
    
    TF_REF_PATTERN = re.compile(r'(\d+)\s*-\s*(?:T|F|TRUE|FALSE|Đ|S|ĐÚNG|SAI)', re.IGNORECASE)
    
    @staticmethod
    def detect_language(text):
        """Detect if document is in English or Vietnamese"""
        vietnamese_chars = set('àáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ')
        vietnamese_count = sum(1 for char in text.lower() if char in vietnamese_chars)
        return 'vi' if vietnamese_count > len(text) * 0.05 else 'en'
    
    @staticmethod
    def detect_question_type(prompt, options, answer=None, order=None, language='en'):
        """Detect question type based on content and answer structure."""
        if not options:
            return Question.QuestionType.FILL_IN_BLANK

        prompt_lower = prompt.lower()

        if order:
            return Question.QuestionType.ORDERING

        has_numbered_options = any("number" in opt for opt in options)
        tf_answer_pattern = re.compile(
            r'^\s*\d+\s*-\s*(?:T|F|TRUE|FALSE|Y|N|YES|NO|Đ|S|ĐÚNG|SAI)(?:\s*,\s*\d+\s*-\s*(?:T|F|TRUE|FALSE|Y|N|YES|NO|Đ|S|ĐÚNG|SAI))*\s*$',
            re.IGNORECASE
        )
        has_tf_answer_format = bool(answer and tf_answer_pattern.match(answer.strip()))
        if has_numbered_options and has_tf_answer_format:
            return Question.QuestionType.TRUE_FALSE_EXTENDED

        ordering_keywords = (
            WordQuestionParser.ORDERING_KEYWORDS_EN
            if language == 'en' else WordQuestionParser.ORDERING_KEYWORDS_VI
        )
        if any(keyword in prompt_lower for keyword in ordering_keywords):
            return Question.QuestionType.ORDERING

        tf_keywords = (
            WordQuestionParser.TRUE_FALSE_KEYWORDS_EN
            if language == 'en' else WordQuestionParser.TRUE_FALSE_KEYWORDS_VI
        )
        if any(keyword in prompt_lower for keyword in tf_keywords):
            return Question.QuestionType.TRUE_FALSE_EXTENDED

        return Question.QuestionType.MULTIPLE_CHOICE

    @staticmethod
    def _extract_paragraph_images(paragraph, doc_part):
        """Extract images from a paragraph element."""
        images = []
        if qn is None:
            return images
        try:
            drawings = paragraph._element.findall('.//' + qn('w:drawing'))
            for drawing in drawings:
                blips = drawing.findall('.//' + qn('a:blip'))
                for blip in blips:
                    rId = blip.get(qn('r:embed'))
                    if rId and rId in doc_part.rels:
                        rel = doc_part.rels[rId]
                        try:
                            image_part = rel.target_part
                            images.append({
                                'blob': image_part.blob,
                                'content_type': image_part.content_type,
                            })
                        except Exception:
                            pass
        except Exception as e:
            print(f"[Parser] Error extracting images from paragraph: {e}")
        return images

    @staticmethod
    def parse_document(file_path):
        """
        Parse Word document and extract raw question data.
        Returns: (list of raw question dicts, language string)
        """
        doc = docx.Document(file_path)
        
        # Detect language
        sample_text = ' '.join([p.text for p in doc.paragraphs[:10]])
        language = WordQuestionParser.detect_language(sample_text)
        print(f"[Parser] Detected language: {language}")
        
        # Select patterns
        if language == 'en':
            question_pattern = WordQuestionParser.QUESTION_PATTERN_EN
            answer_pattern = WordQuestionParser.ANSWER_PATTERN_EN
            order_pattern = WordQuestionParser.ORDER_PATTERN_EN
        else:
            question_pattern = WordQuestionParser.QUESTION_PATTERN_VI
            answer_pattern = WordQuestionParser.ANSWER_PATTERN_VI
            order_pattern = WordQuestionParser.ORDER_PATTERN_VI
        
        raw_questions = []
        current_question = None
        current_options = []
        current_answer = None
        current_order = None
        current_image = None
        
        doc_part = doc.part
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # Check for images in this paragraph
            para_images = WordQuestionParser._extract_paragraph_images(para, doc_part)
            
            if not text:
                # Image-only paragraph: associate with current question
                if para_images and current_question and current_image is None:
                    current_image = para_images[0]
                continue
            
            # If paragraph has text + image, capture the image
            if para_images and current_question and current_image is None:
                current_image = para_images[0]
            
            # Process each line in the paragraph
            lines = text.split('\n')
            for line_text in lines:
                line_text = line_text.strip()
                if not line_text:
                    continue
                
                # Check for question start
                question_match = question_pattern.match(line_text)
                if question_match:
                    # Save previous question
                    if current_question:
                        raw_questions.append({
                            'number': current_question['number'],
                            'prompt': current_question['prompt'],
                            'options': current_options,
                            'answer': current_answer,
                            'order': current_order,
                            'image': current_image,
                        })
                    
                    current_question = {
                        'number': int(question_match.group(1)),
                        'prompt': question_match.group(2).strip(),
                    }
                    current_options = []
                    current_answer = None
                    current_order = None
                    current_image = None
                    continue
                
                # Check for lettered options (A, B, C, D...)
                option_match = WordQuestionParser.OPTION_PATTERN.match(line_text)
                if option_match and current_question:
                    current_options.append({
                        "letter": option_match.group(1),
                        "text": option_match.group(2).strip(),
                        "order": len(current_options)
                    })
                    continue
                
                # Check for numbered options (1, 2, 3... for True/False)
                numbered_match = WordQuestionParser.NUMBERED_OPTION_PATTERN.match(line_text)
                if numbered_match and current_question:
                    current_options.append({
                        "number": int(numbered_match.group(1)),
                        "text": numbered_match.group(2).strip(),
                        "order": len(current_options)
                    })
                    continue
                
                # Check for answer
                answer_match = answer_pattern.match(line_text)
                if answer_match:
                    current_answer = answer_match.group(1).strip()
                    continue
                
                # Check for ordering answer
                order_match = order_pattern.match(line_text)
                if order_match:
                    current_order = order_match.group(1).strip()
                    continue
        
        # Save last question
        if current_question:
            raw_questions.append({
                'number': current_question['number'],
                'prompt': current_question['prompt'],
                'options': current_options,
                'answer': current_answer,
                'order': current_order,
                'image': current_image,
            })
        
        print(f"[Parser] Total raw questions parsed: {len(raw_questions)}")
        return raw_questions, language

    @staticmethod
    def validate_raw_questions(raw_questions, language='en'):
        """
        Validate parsed questions. Returns (valid_questions, errors).
        Questions with errors are excluded from valid_questions.
        """
        errors = []
        valid_questions = []
        
        for q in raw_questions:
            q_num = q['number']
            q_errors = []
            
            # 1. Empty prompt
            if not q['prompt'] or not q['prompt'].strip():
                q_errors.append({
                    'question': q_num,
                    'type': 'empty_prompt',
                    'message': f"Question {q_num}: Empty question prompt"
                })
            
            has_letter_options = any('letter' in opt for opt in q['options'])
            has_numbered_options = any('number' in opt for opt in q['options'])
            
            # 2. Has options but no answer at all
            if (has_letter_options or has_numbered_options) and not q['answer'] and not q['order']:
                q_errors.append({
                    'question': q_num,
                    'type': 'missing_answer',
                    'message': f"Question {q_num}: Missing answer line (ANSWER/ĐÁP ÁN)"
                })
            
            # 3. MC answer references non-existent options
            if q['answer'] and has_letter_options:
                available_letters = {opt['letter'] for opt in q['options'] if 'letter' in opt}
                answer_parts = [l.strip().upper() for l in q['answer'].split(',')]
                is_tf_format = any('-' in part for part in answer_parts)
                
                if not is_tf_format:
                    invalid_refs = [l for l in answer_parts if l and l not in available_letters]
                    if invalid_refs:
                        q_errors.append({
                            'question': q_num,
                            'type': 'invalid_answer_reference',
                            'message': f"Question {q_num}: Answer references non-existent option(s) "
                                       f"{', '.join(invalid_refs)} "
                                       f"(available: {', '.join(sorted(available_letters))})"
                        })
            
            # 4. Ordering references non-existent options
            if q['order'] and has_letter_options:
                available_letters = {opt['letter'] for opt in q['options'] if 'letter' in opt}
                order_letters = [l.strip().upper() for l in q['order'].split(',')]
                invalid_refs = [l for l in order_letters if l and l not in available_letters]
                if invalid_refs:
                    q_errors.append({
                        'question': q_num,
                        'type': 'invalid_order_reference',
                        'message': f"Question {q_num}: Order references non-existent option(s) "
                                   f"{', '.join(invalid_refs)} "
                                   f"(available: {', '.join(sorted(available_letters))})"
                    })
            
            # 5. TF answer references non-existent statement numbers
            if q['answer'] and has_numbered_options:
                available_numbers = {opt['number'] for opt in q['options'] if 'number' in opt}
                tf_matches = WordQuestionParser.TF_REF_PATTERN.findall(q['answer'])
                if tf_matches:
                    referenced_numbers = {int(n) for n in tf_matches}
                    invalid_nums = referenced_numbers - available_numbers
                    if invalid_nums:
                        q_errors.append({
                            'question': q_num,
                            'type': 'invalid_tf_reference',
                            'message': f"Question {q_num}: Answer references non-existent statement(s) "
                                       f"{', '.join(map(str, sorted(invalid_nums)))} "
                                       f"(available: {', '.join(map(str, sorted(available_numbers)))})"
                        })
            
            if q_errors:
                errors.extend(q_errors)
            else:
                valid_questions.append(q)
        
        return valid_questions, errors
    
    @staticmethod
    def _build_question_data(question_info, options, answer, order, language='en'):
        """Build complete question data structure"""
        question_type = WordQuestionParser.detect_question_type(
            question_info["prompt"], options, answer=answer, order=order, language=language
        )
        
        result = {
            "question": {
                "prompt": question_info["prompt"],
                "question_type": question_type,
                "points": 1.0,
                "correct_answer_text": None
            },
            "options": []
        }
        
        if question_type == Question.QuestionType.FILL_IN_BLANK:
            if answer:
                answer_list = [a.strip() for a in answer.split(',') if a.strip()]
                result["question"]["correct_answer_text"] = ", ".join(answer_list)
            else:
                result["question"]["correct_answer_text"] = ""
        
        elif question_type == Question.QuestionType.TRUE_FALSE_EXTENDED:
            answer_map = WordQuestionParser._parse_true_false_answer(answer, language)
            for opt in options:
                opt_number = opt.get("number")
                is_correct = answer_map.get(opt_number, False)
                result["options"].append({
                    "text": opt["text"],
                    "is_correct_bool": is_correct,
                    "correct_order": None,
                    "order": opt["order"]
                })
        
        elif question_type == Question.QuestionType.ORDERING:
            correct_order = WordQuestionParser._parse_ordering_answer(order or answer)
            for opt in options:
                letter = opt.get("letter")
                correct_position = correct_order.index(letter) if letter in correct_order else -1
                result["options"].append({
                    "text": opt["text"],
                    "is_correct_bool": False,
                    "correct_order": correct_position + 1 if correct_position >= 0 else None,
                    "order": opt["order"]
                })
        
        else:  # MULTIPLE_CHOICE
            correct_letters = WordQuestionParser._parse_multiple_choice_answer(answer)
            for opt in options:
                letter = opt.get("letter")
                is_correct = letter in correct_letters
                result["options"].append({
                    "text": opt["text"],
                    "is_correct_bool": is_correct,
                    "correct_order": None,
                    "order": opt["order"]
                })
        
        return result
    
    @staticmethod
    def _parse_true_false_answer(answer, language='en'):
        """Parse True/False answer string"""
        answer_map = {}
        if not answer:
            return answer_map
        
        parts = [p.strip() for p in answer.split(',')]
        for part in parts:
            if '-' in part:
                num_str, tf = part.split('-', 1)
                try:
                    num = int(num_str.strip())
                except ValueError:
                    continue
                tf_upper = tf.strip().upper()
                if language == 'en':
                    is_true = tf_upper in ['T', 'TRUE', 'YES', 'Y']
                else:
                    is_true = tf_upper in ['Đ', 'ĐÚNG', 'T', 'TRUE']
                answer_map[num] = is_true
        
        return answer_map
    
    @staticmethod
    def _parse_ordering_answer(order_str):
        """Parse ordering answer string"""
        if not order_str:
            return []
        return [letter.strip().upper() for letter in order_str.split(',')]
    
    @staticmethod
    def _parse_multiple_choice_answer(answer):
        """Parse multiple choice answer string"""
        if not answer:
            return []
        return [letter.strip().upper() for letter in answer.split(',')]


def process_word_document(file_path, section, user):
    """
    Process Word document and create questions in database.
    With validation, error reporting, and image extraction.
    
    Returns:
        dict with 'success', 'created_count', 'skipped_count', 
        'validation_errors', 'errors', 'language'
    """
    result = {
        "success": False,
        "created_count": 0,
        "skipped_count": 0,
        "skipped_duplicates": [],
        "validation_errors": [],
        "errors": [],
        "language": "unknown"
    }
    
    try:
        print(f"[process_word_document] Starting to process: {file_path}")
        
        # 1. Parse document (raw data + images)
        raw_questions, language = WordQuestionParser.parse_document(file_path)
        result["language"] = language
        
        if not raw_questions:
            result["errors"].append({
                "question": 0,
                "type": "empty_document",
                "message": "No questions found in document. Make sure questions start with "
                           "'Question N:' (English) or 'Câu N:' (Vietnamese)."
            })
            return result
        
        # 2. Validate raw questions
        valid_questions, validation_errors = WordQuestionParser.validate_raw_questions(
            raw_questions, language
        )
        result["validation_errors"] = validation_errors
        
        if validation_errors:
            print(f"[process_word_document] {len(validation_errors)} validation error(s) found")
        
        if not valid_questions:
            result["errors"].append({
                "question": 0,
                "type": "all_invalid",
                "message": "All questions had validation errors. No questions were created."
            })
            return result
        
        # 3. Build fingerprint set for deduplication
        existing_fingerprints = build_existing_fingerprint_set(section)
        
        batch_fingerprints = set()
        
        # 4. Prepare questions for processing
        questions_to_process = []
        skipped_count = 0
        skipped_duplicates = []
        
        for raw_q in valid_questions:
            try:
                # Build structured data
                built = WordQuestionParser._build_question_data(
                    {'number': raw_q['number'], 'prompt': raw_q['prompt']},
                    raw_q['options'],
                    raw_q['answer'],
                    raw_q['order'],
                    language
                )
                
                question_data = built["question"]
                options_data = built["options"]
                
                # Fingerprint check
                new_opts_texts = [opt["text"] for opt in options_data]
                fp = compute_fingerprint(question_data["prompt"], new_opts_texts)
                
                if fp in existing_fingerprints or fp in batch_fingerprints:
                    skipped_count += 1
                    skipped_duplicates.append(question_data["prompt"][:80])
                    continue
                
                batch_fingerprints.add(fp)
                
                questions_to_process.append({
                    'raw_q': raw_q,
                    'question_data': question_data,
                    'options_data': options_data
                })
                
            except Exception as e:
                error_msg = f"Error preparing question {raw_q['number']}: {str(e)}"
                result["errors"].append({
                    "question": raw_q['number'],
                    "type": "preparation_error",
                    "message": error_msg
                })
                print(f"[process_word_document] {error_msg}")
                continue
                
        # 5. Parallel Image Uploads
        image_urls = {}
        images_to_upload = [(idx, q) for idx, q in enumerate(questions_to_process) if q['raw_q'].get('image')]
        
        if images_to_upload:
            print(f"[process_word_document] Uploading {len(images_to_upload)} images in parallel...")
            with ThreadPoolExecutor(max_workers=5) as executor:
                futures = {}
                for idx, q in images_to_upload:
                    image_data = q['raw_q']['image']
                    futures[executor.submit(
                        upload_image_to_cloudinary,
                        image_data['blob'],
                        image_data.get('content_type', 'image/png')
                    )] = idx
                    
                for future in as_completed(futures):
                    idx = futures[future]
                    try:
                        image_urls[idx] = future.result()
                    except Exception as e:
                        print(f"[process_word_document] Image upload failed for index {idx}: {e}")
                        image_urls[idx] = None

        # 6. Bulk Create DB Records
        created_count = 0
        if questions_to_process:
            try:
                with transaction.atomic():
                    questions_to_create = []
                    for idx, q_info in enumerate(questions_to_process):
                        q_data = q_info['question_data']
                        q = Question(
                            section=section,
                            created_by=user,
                            prompt=q_data["prompt"],
                            question_type=q_data["question_type"],
                            correct_answer_text=q_data.get("correct_answer_text"),
                            image=image_urls.get(idx)
                        )
                        questions_to_create.append(q)
                    
                    # Bulk create questions (Returns IDs in PostgreSQL/SQLite)
                    created_questions = Question.objects.bulk_create(questions_to_create)
                    
                    # Prepare options
                    options_to_create = []
                    for created_q, q_info in zip(created_questions, questions_to_process):
                        for opt_data in q_info['options_data']:
                            opt = AnswerOption(
                                question=created_q,
                                text=opt_data["text"],
                                is_correct_bool=opt_data.get("is_correct_bool", False),
                                correct_order=opt_data.get("correct_order"),
                                order=opt_data["order"]
                            )
                            options_to_create.append(opt)
                    
                    # Bulk create options
                    AnswerOption.objects.bulk_create(options_to_create)
                    created_count = len(created_questions)
            except Exception as e:
                error_msg = f"Database transaction failed during bulk creation: {str(e)}"
                result["errors"].append({
                    "question": 0,
                    "type": "database_error",
                    "message": error_msg
                })
                print(f"[process_word_document] {error_msg}")
        
        result["created_count"] = created_count
        result["skipped_count"] = skipped_count
        result["skipped_duplicates"] = skipped_duplicates
        result["success"] = created_count > 0 or skipped_count > 0
        
        print(f"[process_word_document] Created {created_count}, Skipped {skipped_count}, "
              f"Validation errors {len(validation_errors)}")
        
    except Exception as e:
        error_msg = f"Critical error processing document: {str(e)}"
        result["errors"].append({
            "question": 0,
            "type": "critical",
            "message": error_msg
        })
        print(f"[process_word_document] {error_msg}")
        import traceback
        traceback.print_exc()
    
    return result
